'''
Created on 12-May-2018

@author: srinivasan
'''

import logging
import os
import tempfile
from time import strftime, gmtime
import unicodedata

from docx import Document
from unidecode import unidecode_expect_nonascii

from Data_scuff.decorator import singleton
from Data_scuff.utils.PandasUtils import PandasUtils
from lxml.html import fromstring
import sys

logger = logging.getLogger(__name__)


class Utils:
    
    """
    Convert the response into data frame
    """

    @staticmethod
    def convert_pdf_todf(response):
        obj = Utils.createObj()
        __tempPath = obj.__saveFile(response, ".pdf").name
        logging.log('LOG:=======================temp file path: ' + __tempPath)
        dfs = PdfHelper.covertDf(__tempPath).pdfConverter()
        ddf = [PandasUtils.clean_col_names(df) 
               for df in dfs]
        # silent delete the temp file
        obj.__silentremove(__tempPath)
        return ddf
        
    """
    Save response into temp file
    """

    def __saveFile(self, response, suffix):
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp:
            temp.write(response.body)
        return temp
    
    """
    silent delete file in path
    """

    def __silentremove(self, filename):
        import  errno
        try:
            os.remove(filename)
        except OSError as e: 
            if e.errno != errno.ENOENT: 
                raise
    
    @classmethod
    def createObj(cls):
        return cls()
    
    """
    validate the url
    """

    @staticmethod
    def is_valid_url(url):
        try:
            import re
            regex = re.compile(
                r'^(?:http)s?://'  
                r'(?:(?:[A-Z0-9](?:[A-Z0-9-]{0,61}[A-Z0-9])?\.)+'
                r'(?:[A-Z]{2,6}\.?|[A-Z0-9-]{2,}\.?)|'  
                r'localhost|'  
                r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})' 
                r'(?::\d+)?' 
                r'(?:/?|[/?]\S+)$', re.I)
            if regex.match(url) or url == 'about:blank' or url == 'data:,':
                return True
            else:
                return False 
        except Exception:
            return False
    
    """
    extracting domain from the URL
    """

    @staticmethod
    def get_domain_url(url):
        url_header = url.split('://')[0]
        simple_url = url.split('://')[1]
        base_url = simple_url.split('/')[0]
        domain_url = url_header + '://' + base_url
        return domain_url
    
    @staticmethod
    def getingestion_timestamp():
        return strftime("%Y%m%d_%H%M%S", gmtime())
    
    @staticmethod
    def getRundateFileName(name, version='v1'):
        return '{}_{}_{}.csv'.format(name, strftime("%Y%m%d", gmtime()), version)
    
    @staticmethod
    def getCustomFileName(*args):
        return '{0}_{1}_{2}.csv'.format(*args)
    
    @staticmethod
    def getDateStr(_dt, _format='%d-%b-%y'):
        if _dt:
            from dateutil import parser
            dt = parser.parse(_dt)
            return dt.strftime(_format)
        return ''
    
    @staticmethod
    def removeUnicode(val):
        if val is None or val == '':
            return val
        return unicodedata.normalize("NFKD", val)
    
    @staticmethod
    def removeNonAscii(val):
        return unidecode_expect_nonascii(val).replace(",ai", "-")
    
    @staticmethod
    def getUsername():
        import getpass
        return getpass.getuser()
    
    @staticmethod
    def getCurrentDateTimeStr():
        return strftime("%Y-%b-%d %H:%M:%S", gmtime())
    
    @staticmethod
    def parse_form(html):
        tree = fromstring(html)
        return [{e.get('name') or e.get('id') or e.get('class'):
                  {l.get('name'):l.get('value') for l in e.cssselect('input') 
                   if l.get('name')}} for e in tree.cssselect('form')]
    
    def get_size(self, obj, seen=None):
        size = sys.getsizeof(obj)
        if seen is None:
            seen = set()
        obj_id = id(obj)
        if obj_id in seen:
            return 0
        seen.add(obj_id)
        if isinstance(obj, dict):
            size += sum([self.get_size(v, seen) for v in obj.values()])
            size += sum([self.get_size(k, seen) for k in obj.keys()])
        elif hasattr(obj, '__dict__'):
            size += self.get_size(obj.__dict__, seen)
        elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
            size += sum([self.get_size(i, seen) for i in obj])
        return size
    
    @staticmethod
    def getExcelHeaders(requirement_path):
        headers = {}
        import pandas as pd
        data = pd.read_excel(requirement_path, sheet_name=0)
        my_map = data.head(1).fillna('').to_dict(orient='records')[0]
        null_header = [' '.join(key.replace('\n', ' ').replace('\r', '').lower().split())
                        for key, value in my_map.items() if not value]
        headers['null_header'] = null_header
        d = {' '.join(key.replace('\n', ' ').replace('\r', '').split())
             : ' '.join(value.replace('\n', '_').replace('\r', '').split())
             if value else key.replace('\n', '_').replace('\r', '').lower() 
              for key, value in my_map.items()}
        inv_map = {v:'' if 'Unnamed' in k else k for k, v in d.items()}
        headers['feed_expo'] = [*inv_map]
        headers['top_header'] = inv_map
        return headers


class DocXHelper:
    
    def readTables(self, path, isheader=True):
        wordDoc = Document(path)
        return [self.__readTable(table.rows, isheader) \
                for table in wordDoc.tables]

    """
    get Header from the first row
    """

    def __getHeader(self, row):
        return [cell.text.strip() for cell in row.cells]
    
    def readTable(self, path, isheader=True):
        rows = Document(path).tables[0].rows
        return self.__readTable(rows, isheader)
    
    """
    Read table convert into data frame
    """

    def __readTable(self, rows, isheader):
        header = self.__getHeader(rows.pop(0)) if isheader \
         else [ i for i in range(0, rows[0].height)]
        return PandasUtils.convertlistDictToDf([{h:cell.text for (h, cell) \
                                                 in zip(row.cells, header)} for row in rows])


@singleton
class ExtractPDFTables:
    
    def __getValue(self, t):
        value = t.text
        if value is None:
            value = t.find('b').text
        return value
    
    def extractTableData(self, url, options=""):
        import xml.etree.ElementTree as ET
        if isinstance(url, str):
            from urllib.request import urlopen
            html = urlopen(url).read()
        else:
            html = url.body
        tree = ET.fromstring(self.pdftoxml(html, options))
        xml_pages = tree.findall(".//page")
        values = []
        for p in xml_pages:
            texts = p.findall("text")
            top = int(texts[0].get('top'))
            prev_left = 0
            row = []
            rows = []
            cell = []
            for t in texts:
                current_left = int(t.get('left'))
                if top == int(t.get('top')) or prev_left == current_left:
                    value = self.__getValue(t)
                    if prev_left == 0 or prev_left == current_left:
                        cell.append('' if value is None else value.strip())
                    else:
                        row.append(" ".join(cell))
                        cell.clear()
                        cell.append('' if value is None else value.strip())
                else:
                    value = self.__getValue(t)
                    top = int(t.get('top'))
                    row.append(" ".join(cell))
                    rows.append(row.copy())
                    row.clear()
                    cell.clear()
                    cell.append('' if value is None else value)
                prev_left = current_left
            values.append(rows)
        return values
    
    """converts pdf file to xml file"""

    def pdftoxml(self, pdfdata, options=""):
        pdffout = tempfile.NamedTemporaryFile(suffix='.pdf')
        pdffout.write(pdfdata)
        pdffout.flush()
        xmlin = tempfile.NamedTemporaryFile(mode='r', suffix='.xml')
        tmpxml = xmlin.name
        cmd = 'pdftohtml -xml -nodrm -zoom 1.5 -enc UTF-8 -noframes %s "%s" "%s"' % (
            options, pdffout.name, os.path.splitext(tmpxml)[0])
        cmd = cmd + " >/dev/null 2>&1"
        os.system(cmd)
        pdffout.close()
        xmldata = xmlin.read()
        xmlin.close()
        return xmldata

    
class PdfHelper(object):
    
    def __init__(self, path):
        self.path = path
        
    @classmethod
    def covertDf(cls, path):
        return cls(path)
    
    def pdfConverter(self):
        from tabula.wrapper import read_pdf
        return read_pdf(self.path, pages='all', lattice=True)
